"""
word_generator.py — Genere un PV en format Word (.docx) a partir de l analyse AG
Structure professionnelle : en-tete, tableau participants, resolutions, signatures, disclaimer.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _ajouter_ligne_tableau(tableau, cellules: list, gras: bool = False, fond: str = None):
    """Ajoute une ligne dans un tableau Word avec style optionnel."""
    ligne = tableau.add_row()
    for i, texte in enumerate(cellules):
        cellule = ligne.cells[i]
        cellule.text = str(texte) if texte is not None else ""
        para = cellule.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(cellule.text)
        run.bold = gras
        if fond:
            tcPr = cellule._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fond)
            tcPr.append(shd)
    return ligne


def generer_pv_word(analyse: dict, chemin_sortie: str) -> str:
    """
    Genere un PV d AG au format Word (.docx) a partir de l analyse structuree.

    Args:
        analyse: JSON structure produit par analyzer.analyser_transcription()
        chemin_sortie: Chemin du fichier .docx a creer

    Returns:
        str: Chemin du fichier cree
    """
    doc = Document()

    # ── Configuration des marges ──────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # ── Extraction donnees format v2 ──────────────────────────────────────────
    infos = analyse.get("informations_generales", {})
    entite = infos.get("entite", analyse.get("entite", "Entite inconnue"))
    date_ag = infos.get("date", analyse.get("date", "date inconnue"))
    lieu = infos.get("lieu", analyse.get("lieu", "lieu inconnu"))
    type_asm = infos.get("type_assemblee", "AG")
    heure_ouv = infos.get("heure_ouverture", "")
    heure_clo = infos.get("heure_cloture", "")
    president = infos.get("president_seance", "")
    secretaire = infos.get("secretaire", "")

    participants = analyse.get("participants", {})
    type_ag = analyse.get("type_ag", "autre")
    resolutions = analyse.get("resolutions", [])
    points_divers = analyse.get("points_divers", analyse.get("incidents_divers", []))
    decisions = analyse.get("decisions_finales", [])
    conformite = analyse.get("conformite_legale", {})
    meta = analyse.get("meta", {})

    # ── EN-TETE ───────────────────────────────────────────────────────────────
    titre_para = doc.add_paragraph()
    titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre_para.add_run("PROCES-VERBAL D ASSEMBLEE GENERALE")
    run_titre.bold = True
    run_titre.font.size = Pt(16)

    sous_titre_para = doc.add_paragraph()
    sous_titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sous = sous_titre_para.add_run(f"{type_asm.upper()}")
    run_sous.bold = True
    run_sous.font.size = Pt(13)

    doc.add_paragraph()

    # ── INFORMATIONS GENERALES ────────────────────────────────────────────────
    doc.add_heading("1. Informations generales", level=1)

    table_infos = doc.add_table(rows=0, cols=2)
    table_infos.style = "Table Grid"
    lignes_infos = [
        ("Entite / Denomination", entite),
        ("Date de l assemblee", date_ag),
        ("Lieu", lieu),
        ("Type d assemblee", type_asm),
        ("Heure d ouverture", heure_ouv or "non mentionne"),
        ("Heure de cloture", heure_clo or "non mentionne"),
        ("President de seance", president or "non mentionne"),
        ("Secretaire de seance", secretaire or "non mentionne"),
    ]
    for label, valeur in lignes_infos:
        ligne = table_infos.add_row()
        ligne.cells[0].text = label
        ligne.cells[0].paragraphs[0].runs[0].bold = True
        ligne.cells[1].text = str(valeur)

    doc.add_paragraph()

    # ── PARTICIPANTS ET QUORUM ────────────────────────────────────────────────
    doc.add_heading("2. Participants et quorum", level=1)

    total_presents = participants.get("total_presents", "N/A")
    if isinstance(total_presents, list):
        total_presents = len(total_presents)
    total_rep = participants.get("total_representes", "N/A")
    if isinstance(total_rep, list):
        total_rep = len(total_rep)
    total_votants = participants.get("total_votants", "N/A")
    total_voix = participants.get("total_voix", "N/A")
    quorum_atteint = participants.get("quorum_atteint")
    quorum_requis = participants.get("quorum_requis", "")

    table_quorum = doc.add_table(rows=0, cols=2)
    table_quorum.style = "Table Grid"
    lignes_q = [
        ("Membres / associes presents", str(total_presents)),
        ("Membres representes (procurations)", str(total_rep)),
        ("Total votants", str(total_votants)),
        ("Total voix / parts", str(total_voix)),
        ("Quorum requis", quorum_requis or "Selon statuts"),
        ("Quorum atteint", "OUI" if quorum_atteint is True else "NON" if quorum_atteint is False else "N/A"),
    ]
    for label, valeur in lignes_q:
        ligne = table_quorum.add_row()
        ligne.cells[0].text = label
        ligne.cells[0].paragraphs[0].runs[0].bold = True
        cell_val = ligne.cells[1]
        cell_val.text = valeur
        if label == "Quorum atteint":
            run = cell_val.paragraphs[0].runs[0]
            run.bold = True
            if quorum_atteint is True:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif quorum_atteint is False:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    # ── ORDRE DU JOUR ─────────────────────────────────────────────────────────
    odj = analyse.get("ordre_du_jour", [])
    if odj:
        doc.add_heading("3. Ordre du jour", level=1)
        for pt in odj:
            p = doc.add_paragraph(style="List Number")
            p.add_run(pt.get("intitule", ""))
        doc.add_paragraph()

    # ── RESOLUTIONS ───────────────────────────────────────────────────────────
    doc.add_heading("4. Resolutions", level=1)

    if not resolutions:
        doc.add_paragraph("Aucune resolution identifiee.")
    else:
        for r in resolutions:
            statut = r.get("statut", r.get("resultat", ""))
            is_adopte = statut in ("adoptee", "adoptee", "adopte")
            titre_res = r.get("titre", r.get("intitule", f"Resolution {r.get('numero', '?')}"))

            # Titre resolution
            h = doc.add_heading(f"Resolution n°{r.get('numero', '?')} - {titre_res}", level=2)

            # Description
            if r.get("description"):
                doc.add_paragraph(r["description"])

            # Discussion
            discussions = r.get("points_discussion", [])
            if discussions:
                p_disc = doc.add_paragraph()
                p_disc.add_run("Discussion : ").bold = True
                p_disc.add_run(" / ".join(discussions))

            # Tableau de votes
            votes = r.get("votes", {})
            unite = votes.get("unite", "voix")
            abstentions = votes.get("abstentions", votes.get("abstention", "N/A"))

            table_votes = doc.add_table(rows=0, cols=4)
            table_votes.style = "Table Grid"
            # En-tete
            en_tete = table_votes.add_row()
            for i, texte in enumerate(["Pour", "Contre", "Abstentions", "Unite"]):
                en_tete.cells[i].text = texte
                en_tete.cells[i].paragraphs[0].runs[0].bold = True
                tcPr = en_tete.cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E1F2")
                tcPr.append(shd)
            # Donnees
            data_row = table_votes.add_row()
            data_row.cells[0].text = str(votes.get("pour", "N/A"))
            data_row.cells[1].text = str(votes.get("contre", "N/A"))
            data_row.cells[2].text = str(abstentions)
            data_row.cells[3].text = unite

            # Resultat
            p_res = doc.add_paragraph()
            run_label = p_res.add_run("Resultat : ")
            run_label.bold = True
            run_val = p_res.add_run(statut.upper() if statut else "?")
            run_val.bold = True
            if statut in ("adoptee", "adoptee", "adopte"):
                run_val.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif statut in ("rejetee", "rejetee", "rejete"):
                run_val.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            if r.get("base_legale") and r["base_legale"] != "non mentionne dans la transcription":
                p_bl = doc.add_paragraph()
                p_bl.add_run("Base legale : ").bold = True
                p_bl.add_run(r["base_legale"])

            doc.add_paragraph()

    # ── DECISIONS FINALES ─────────────────────────────────────────────────────
    if decisions:
        doc.add_heading("5. Recapitulatif des decisions", level=1)
        table_dec = doc.add_table(rows=0, cols=3)
        table_dec.style = "Table Grid"
        # En-tete
        en_tete = table_dec.add_row()
        for i, texte in enumerate(["Resolution", "Titre", "Statut"]):
            en_tete.cells[i].text = texte
            en_tete.cells[i].paragraphs[0].runs[0].bold = True
        for d in decisions:
            ligne = table_dec.add_row()
            ligne.cells[0].text = str(d.get("resolution", ""))
            ligne.cells[1].text = d.get("titre", "")
            st = d.get("statut", "")
            ligne.cells[2].text = st.upper() if st else ""
        doc.add_paragraph()

    # ── QUESTIONS DIVERSES ────────────────────────────────────────────────────
    if points_divers:
        doc.add_heading("6. Questions diverses", level=1)
        for pt in points_divers:
            doc.add_paragraph(pt, style="List Bullet")
        doc.add_paragraph()

    # ── CONFORMITE LEGALE ─────────────────────────────────────────────────────
    alertes = conformite.get("alertes", [])
    recommandations = conformite.get("recommandations", [])
    if alertes or recommandations:
        doc.add_heading("7. Points de conformite legale", level=1)
        if alertes:
            p = doc.add_paragraph()
            p.add_run("Alertes :").bold = True
            for al in alertes:
                doc.add_paragraph(al, style="List Bullet")
        if recommandations:
            p = doc.add_paragraph()
            p.add_run("Recommandations :").bold = True
            for rec in recommandations:
                doc.add_paragraph(rec, style="List Bullet")
        doc.add_paragraph()

    # ── SIGNATURES ────────────────────────────────────────────────────────────
    doc.add_heading("Signatures", level=1)

    table_sig = doc.add_table(rows=3, cols=2)
    table_sig.style = "Table Grid"
    table_sig.rows[0].cells[0].text = "Le President de seance"
    table_sig.rows[0].cells[1].text = "Le Secretaire de seance"
    table_sig.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table_sig.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    table_sig.rows[1].cells[0].text = president or "___________________________"
    table_sig.rows[1].cells[1].text = secretaire or "___________________________"
    table_sig.rows[2].cells[0].text = "Date et signature :"
    table_sig.rows[2].cells[1].text = "Date et signature :"

    doc.add_paragraph()

    # ── DISCLAIMER ────────────────────────────────────────────────────────────
    p_dis = doc.add_paragraph()
    p_dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dis = p_dis.add_run(
        "Document genere par AG Assistant avec assistance IA | "
        f"Genere le {datetime.now().strftime('%d/%m/%Y a %H:%M')} | "
        "Document PROVISOIRE - doit etre relu, corrige et signe avant toute valeur juridique."
    )
    run_dis.italic = True
    run_dis.font.size = Pt(8)
    run_dis.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if meta.get("avertissement"):
        p_warn = doc.add_paragraph()
        p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_warn = p_warn.add_run(meta["avertissement"])
        run_warn.italic = True
        run_warn.font.size = Pt(8)
        run_warn.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

    doc.save(chemin_sortie)
    return chemin_sortie
