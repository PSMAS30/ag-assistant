"""
test_convocation_et_presence.py — Tests convocation_generator et presence_generator
Sans appel API, sans audio.
"""

import json
import os
import tempfile
import pytest
from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures"

def _fixture(nom):
    with open(FIXTURES_DIR / nom, encoding="utf-8") as f:
        return json.load(f)


# ── Convocation ──────────────────────────────────────────────────────────────

def test_convocation_pdf_cree_fichier():
    import convocation_generator
    fixture = _fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        chemin = tmp.name
    try:
        convocation_generator.generer_convocation_pdf(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 1000
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_convocation_word_cree_fichier():
    import convocation_generator
    fixture = _fixture("analyse_association.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        convocation_generator.generer_convocation_word(fixture, chemin, date_ag_proposee="15/09/2025")
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 5000
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_convocation_word_contient_odj():
    from docx import Document
    import convocation_generator
    fixture = _fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        convocation_generator.generer_convocation_word(fixture, chemin)
        doc = Document(chemin)
        texte = " ".join(p.text for p in doc.paragraphs)
        assert "Ordre du jour" in texte or "ordre du jour" in texte.lower()
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_convocation_dict_vide_ne_plante_pas():
    import convocation_generator
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        chemin = tmp.name
    try:
        convocation_generator.generer_convocation_pdf({}, chemin)
        assert os.path.exists(chemin)
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


# ── Feuille de presence ──────────────────────────────────────────────────────

def test_presence_pdf_cree_fichier():
    import presence_generator
    fixture = _fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        chemin = tmp.name
    try:
        presence_generator.generer_feuille_presence_pdf(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 1000
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_presence_word_cree_fichier():
    import presence_generator
    fixture = _fixture("analyse_pme.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        presence_generator.generer_feuille_presence_word(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 5000
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_presence_word_contient_entite():
    from docx import Document
    import presence_generator
    fixture = _fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        presence_generator.generer_feuille_presence_word(fixture, chemin)
        doc = Document(chemin)
        textes = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    textes.append(cell.text)
        texte_complet = " ".join(textes)
        assert "Acacias" in texte_complet
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


def test_presence_dict_vide_ne_plante_pas():
    import presence_generator
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        presence_generator.generer_feuille_presence_word({}, chemin)
        assert os.path.exists(chemin)
    finally:
        if os.path.exists(chemin): os.unlink(chemin)


# ── Comparaison N vs N-1 ─────────────────────────────────────────────────────

def test_comparer_ag(tmp_path, monkeypatch):
    import historique_manager
    monkeypatch.setattr(historique_manager, "HISTORIQUE_DIR", tmp_path)
    f1 = _fixture("analyse_copropriete.json")
    f2 = _fixture("analyse_copropriete.json")
    # Modifier legèrement f2 pour simuler une evolution
    f2["participants"]["total_votants"] = 20
    c1 = historique_manager.sauvegarder_ag(f1)
    c2 = historique_manager.sauvegarder_ag(f2)
    rapport = historique_manager.comparer_ag(c1, c2)
    assert "resolutions_communes" in rapport
    assert "quorum" in rapport
    assert "participants" in rapport


def test_export_import_historique(tmp_path, monkeypatch):
    import historique_manager
    monkeypatch.setattr(historique_manager, "HISTORIQUE_DIR", tmp_path)
    fixture = _fixture("analyse_association.json")
    historique_manager.sauvegarder_ag(fixture)
    zip_bytes = historique_manager.exporter_historique()
    assert len(zip_bytes) > 0
    # Importer dans un nouveau dossier
    tmp2 = tmp_path / "import"
    tmp2.mkdir()
    monkeypatch.setattr(historique_manager, "HISTORIQUE_DIR", tmp2)
    nb = historique_manager.importer_historique(zip_bytes)
    assert nb == 1
    assert historique_manager.nb_ag_sauvegardees() == 1
