"""
test_word_generator.py — Tests word_generator.py (aucun appel API reel)
Verifie la generation de fichiers .docx a partir des fixtures.
"""

import json
import pytest
import tempfile
import os
from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures"


def _charger_fixture(nom: str) -> dict:
    with open(FIXTURES_DIR / nom, encoding="utf-8") as f:
        return json.load(f)


def test_word_copropriete_cree_fichier():
    """generer_pv_word cree un fichier .docx pour la fixture copropriete."""
    import word_generator
    fixture = _charger_fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 5000
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)


def test_word_association_cree_fichier():
    """generer_pv_word cree un fichier .docx pour la fixture association."""
    import word_generator
    fixture = _charger_fixture("analyse_association.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 5000
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)


def test_word_pme_cree_fichier():
    """generer_pv_word cree un fichier .docx pour la fixture PME."""
    import word_generator
    fixture = _charger_fixture("analyse_pme.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word(fixture, chemin)
        assert os.path.exists(chemin)
        assert os.path.getsize(chemin) > 5000
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)


def test_word_contenu_lisible():
    """Le .docx genere contient bien les donnees de l entite."""
    from docx import Document
    import word_generator
    fixture = _charger_fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word(fixture, chemin)
        doc = Document(chemin)
        # Lire paragraphes ET cellules de tableaux (les infos sont dans des tables)
        textes = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    textes.append(cell.text)
        texte_complet = " ".join(textes)
        assert "Acacias" in texte_complet
        assert "Lefebvre" in texte_complet
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)


def test_word_contient_resolutions():
    """Le .docx contient des titres de resolutions."""
    from docx import Document
    import word_generator
    fixture = _charger_fixture("analyse_copropriete.json")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word(fixture, chemin)
        doc = Document(chemin)
        texte_complet = " ".join(p.text for p in doc.paragraphs)
        assert "Resolution" in texte_complet or "resolution" in texte_complet.lower()
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)


def test_word_analyse_vide_ne_plante_pas():
    """generer_pv_word ne plante pas avec un dict vide."""
    import word_generator
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        chemin = tmp.name
    try:
        word_generator.generer_pv_word({}, chemin)
        assert os.path.exists(chemin)
    finally:
        if os.path.exists(chemin):
            os.unlink(chemin)
