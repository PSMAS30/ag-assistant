"""
convocation_generator.py — Genere un document de convocation conforme par type d entite
Export PDF et Word. Mentions legales obligatoires pre-remplies.
"""

from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REGLES_DELAIS = {
    "copropriete": "21 jours minimum avant la tenue de l AG (art. 9 decret 17 mars 1967)",
    "association": "Selon les statuts (generalement 15 jours minimum)",
    "pme_sas": "Selon les statuts SAS",
    "pme_sarl": "15 jours minimum pour les AGO (art. L223-27 Code de commerce)",
    "pme_autre": "Selon les statuts",
    "autre": "Selon les statuts ou la reglementation applicable",
}

MENTIONS_OBLIGATOIRES = {
    "copropriete": [
        "Denomination et adresse de la copropriete",
        "Date, heure et lieu de l assemblee",
        "Ordre du jour complet",
        "Documents joints : comptes, budget, devis",
        "Feuille de presence et formulaire de procuration",
        "Numero de convocation (1ere ou 2eme)",
    ],
    "association": [
        "Denomination et siege de l association",
        "Type d assemblee (ordinaire / extraordinaire)",
        "Date, heure et lieu",
        "Ordre du jour",
        "Modalites de vote (procuration si autorisee par statuts)",
    ],
    "pme_sas": [
        "Denomination sociale, forme juridique, capital, siege social",
        "Date, heure et lieu",
        "Ordre du jour",
        "Modalites de participation et de vote",
        "Documents mis a disposition des associes",
    ],
    "pme_sarl": [
        "Denomination sociale, capital social, siege",
        "Date, heure et lieu",
        "Ordre du jour",
        "Droit de communication des associes (art. L223-26)",
        "Modalites de vote par correspondance si prevues",
    ],
}


def _sanitiser(texte: str) -> str:
    if not texte:
        return ""
    rempl = {"→": "->", "—": "-", "è": "e", "é": "e", "ê": "e", "à": "a", "ù": "u", "ô": "o", "î": "i", "û": "u", "ç": "c"}
    for k, v in rempl.items():
        texte = texte.replace(k, v)
    return texte.encode("latin-1", errors="ignore").decode("latin-1")


def generer_convocation_word(analyse: dict, chemin_sortie: str, date_convocation: str = None, date_ag_proposee: str = None, lieu_propose: str = None) -> str:
    """
    Genere une convocation Word depuis l analyse d une AG precedente ou les infos de base.

    Args:
        analyse: JSON structure (peut etre vide pour une convocation from scratch)
        chemin_sortie: Chemin du fichier .docx
        date_convocation: Date d envoi de la convocation (defaut: aujourd hui)
        date_ag_proposee: Date proposee pour l AG
        lieu_propose: Lieu propose pour l AG

    Returns:
        str: Chemin du fichier cree
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    infos = analyse.get("informations_generales", {})
    entite = infos.get("entite", analyse.get("entite", "Entite"))
    type_ag = analyse.get("type_ag", "autre")
    type_asm = infos.get("type_assemblee", "Assemblee Generale")
    odj = analyse.get("ordre_du_jour", [])
    date_conv = date_convocation or datetime.now().strftime("%d/%m/%Y")
    delai = REGLES_DELAIS.get(type_ag, REGLES_DELAIS["autre"])
    mentions = MENTIONS_OBLIGATOIRES.get(type_ag, MENTIONS_OBLIGATOIRES.get("association", []))

    # En-tete
    p_entite = doc.add_paragraph()
    p_entite.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_e = p_entite.add_run(entite.upper())
    run_e.bold = True
    run_e.font.size = Pt(13)

    doc.add_paragraph()

    # Objet
    p_obj = doc.add_paragraph()
    p_obj.add_run("Objet : ").bold = True
    p_obj.add_run(f"Convocation a l {type_asm}")

    p_date = doc.add_paragraph()
    p_date.add_run(f"Date de convocation : ").bold = True
    p_date.add_run(date_conv)

    doc.add_paragraph()

    # Corps
    doc.add_heading("Madame, Monsieur,", level=2)
    doc.add_paragraph()

    p_intro = doc.add_paragraph(
        f"Vous etes convoque(e) a l {type_asm} de {entite} qui se tiendra :"
    )

    # Date / lieu / heure
    table_details = doc.add_table(rows=0, cols=2)
    table_details.style = "Table Grid"
    for label, val in [
        ("Date", date_ag_proposee or "________________"),
        ("Heure", "________________"),
        ("Lieu", lieu_propose or infos.get("lieu", "________________")),
    ]:
        row = table_details.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val

    doc.add_paragraph()

    # Ordre du jour
    doc.add_heading("Ordre du jour", level=2)
    if odj:
        for pt in odj:
            p = doc.add_paragraph(style="List Number")
            p.add_run(pt.get("intitule", ""))
    else:
        for i in range(1, 5):
            p = doc.add_paragraph(style="List Number")
            p.add_run("________________")

    doc.add_paragraph()

    # Mentions legales
    doc.add_heading("Informations legales", level=2)
    p_delai = doc.add_paragraph()
    p_delai.add_run("Delai de convocation : ").bold = True
    p_delai.add_run(delai)
    doc.add_paragraph()
    p_men = doc.add_paragraph()
    p_men.add_run("Cette convocation contient les mentions obligatoires suivantes :").bold = True
    for mention in mentions:
        doc.add_paragraph(mention, style="List Bullet")

    doc.add_paragraph()

    # Procuration
    doc.add_heading("Procuration (si vous ne pouvez pas assister)", level=2)
    doc.add_paragraph("Si vous ne pouvez pas assister a cette assemblee, vous pouvez vous faire representer par un mandataire de votre choix en lui remettant le formulaire de procuration ci-joint.")
    doc.add_paragraph()

    # Signature
    p_sig = doc.add_paragraph()
    p_sig.add_run(f"Fait a ________________, le {date_conv}")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("Le President / Le Syndic / Le Gerant")
    doc.add_paragraph()
    doc.add_paragraph("Nom : ___________________________")
    doc.add_paragraph("Signature : ___________________________")
    doc.add_paragraph()

    # Disclaimer
    p_dis = doc.add_paragraph()
    p_dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dis = p_dis.add_run(
        f"Convocation generee par AG Assistant — {datetime.now().strftime('%d/%m/%Y')} — Document provisoire a adapter et valider"
    )
    run_dis.italic = True
    run_dis.font.size = Pt(8)
    run_dis.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(chemin_sortie)
    return chemin_sortie


def generer_convocation_pdf(analyse: dict, chemin_sortie: str, date_convocation: str = None, date_ag_proposee: str = None, lieu_propose: str = None) -> str:
    """Genere une convocation PDF."""
    infos = analyse.get("informations_generales", {})
    entite = _sanitiser(infos.get("entite", "Entite"))
    type_ag = analyse.get("type_ag", "autre")
    type_asm = _sanitiser(infos.get("type_assemblee", "Assemblee Generale"))
    odj = analyse.get("ordre_du_jour", [])
    date_conv = date_convocation or datetime.now().strftime("%d/%m/%Y")
    delai = _sanitiser(REGLES_DELAIS.get(type_ag, "Selon statuts"))
    lieu = _sanitiser(lieu_propose or infos.get("lieu", "________________"))

    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    largeur = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(largeur, 8, entite, align="C")
    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(largeur, 6, f"CONVOCATION - {type_asm.upper()}", align="C")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(60, 6, "Date de convocation :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 60, 6, date_conv, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(largeur, 6, f"Vous etes convoque(e) a l {type_asm} qui se tiendra :")
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(30, 6, "Date :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 30, 6, date_ag_proposee or "________________", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(30, 6, "Heure :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 30, 6, "________________", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(30, 6, "Lieu :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 30, 6, lieu[:80], new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(largeur, 6, "ORDRE DU JOUR")
    pdf.set_font("Helvetica", "", 10)
    if odj:
        for i, pt in enumerate(odj, 1):
            pdf.multi_cell(largeur, 6, f"{i}. {_sanitiser(pt.get('intitule', ''))}")
    else:
        for i in range(1, 5):
            pdf.cell(largeur, 6, f"{i}. ________________", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(largeur, 6, "Delai legal de convocation :", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(largeur, 5, delai)
    pdf.ln(6)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur, 6, f"Fait a ________________, le {date_conv}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.cell(largeur // 2, 6, "Le President / Syndic / Gerant :", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.cell(largeur // 2, 6, "Nom : _______________________", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(largeur // 2, 6, "Signature : _______________________", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    pdf.set_font("Helvetica", "I", 8)
    pdf.multi_cell(largeur, 5, f"Convocation generee par AG Assistant le {datetime.now().strftime('%d/%m/%Y')} - Document provisoire a valider", align="C")

    pdf.output(chemin_sortie)
    return chemin_sortie
