"""
presence_generator.py — Genere une feuille de presence pre-remplie depuis l analyse AG
Export PDF et Word. Compatible format JSON v2.
"""

from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _sanitiser(texte: str) -> str:
    """Remplace les caracteres hors latin-1 pour compatibilite PDF."""
    if not texte:
        return ""
    rempl = {"→": "->", "—": "-", "✅": "[OK]", "❌": "[NON]", "⚠️": "[!]"}
    for k, v in rempl.items():
        texte = texte.replace(k, v)
    return texte.encode("latin-1", errors="ignore").decode("latin-1")


def generer_feuille_presence_word(analyse: dict, chemin_sortie: str) -> str:
    """
    Genere une feuille de presence Word pre-remplie depuis l analyse.

    Args:
        analyse: JSON structure produit par analyzer
        chemin_sortie: Chemin du fichier .docx

    Returns:
        str: Chemin du fichier cree
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    infos = analyse.get("informations_generales", {})
    entite = infos.get("entite", "Entite inconnue")
    date_ag = infos.get("date", "")
    lieu = infos.get("lieu", "")
    type_asm = infos.get("type_assemblee", "AG")
    participants = analyse.get("participants", {})
    type_ag = analyse.get("type_ag", "autre")

    # Titre
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titre.add_run("FEUILLE DE PRESENCE")
    run_t.bold = True
    run_t.font.size = Pt(16)

    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sous.add_run(f"{type_asm.upper()} — {entite}").bold = True

    doc.add_paragraph()

    # Informations
    table_info = doc.add_table(rows=0, cols=2)
    table_info.style = "Table Grid"
    for label, val in [
        ("Entite", entite),
        ("Date", date_ag or "________________"),
        ("Lieu", lieu or "________________"),
        ("Type d assemblee", type_asm),
        ("Heure d ouverture", infos.get("heure_ouverture", "________________")),
    ]:
        row = table_info.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(val)

    doc.add_paragraph()

    # Tableau des presents (pre-rempli si disponible)
    presents = participants.get("presents", [])
    representes = participants.get("representes", [])

    # Colonne unite selon type
    if type_ag == "copropriete":
        col_unite = "Tantièmes / Millièmes"
    elif type_ag in ("pme_sas", "pme_sarl", "pme_autre"):
        col_unite = "Parts / Actions"
    else:
        col_unite = "Voix"

    doc.add_heading("Membres presents", level=2)
    nb_lignes_presents = max(len(presents) if isinstance(presents, list) else 0, 10)
    table_p = doc.add_table(rows=1, cols=4)
    table_p.style = "Table Grid"
    entetes = ["Nom / Prenom", col_unite, "Signature", "Observations"]
    for i, e in enumerate(entetes):
        table_p.rows[0].cells[i].text = e
        table_p.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    # Pre-remplir avec les donnees connues
    if isinstance(presents, list):
        for p in presents:
            nom = p.get("nom", "") if isinstance(p, dict) else str(p)
            voix = str(p.get("voix_ou_parts", "")) if isinstance(p, dict) else ""
            row = table_p.add_row()
            row.cells[0].text = nom
            row.cells[1].text = voix
            row.cells[2].text = ""
            row.cells[3].text = p.get("observations", "") if isinstance(p, dict) else ""

    # Lignes vides supplementaires
    lignes_remplies = len(presents) if isinstance(presents, list) else 0
    for _ in range(max(0, 10 - lignes_remplies)):
        table_p.add_row()

    doc.add_paragraph()

    # Representes
    doc.add_heading("Membres representes (procurations)", level=2)
    table_r = doc.add_table(rows=1, cols=4)
    table_r.style = "Table Grid"
    for i, e in enumerate(["Mandant (represente)", "Mandataire (representant)", col_unite, "Signature mandataire"]):
        table_r.rows[0].cells[i].text = e
        table_r.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    if isinstance(representes, list):
        for r in representes:
            mandant = r.get("mandant", "") if isinstance(r, dict) else ""
            mandataire = r.get("mandataire", "") if isinstance(r, dict) else ""
            voix = str(r.get("voix_ou_parts", "")) if isinstance(r, dict) else ""
            row = table_r.add_row()
            row.cells[0].text = mandant
            row.cells[1].text = mandataire
            row.cells[2].text = voix
            row.cells[3].text = ""

    for _ in range(max(0, 5 - (len(representes) if isinstance(representes, list) else 0))):
        table_r.add_row()

    doc.add_paragraph()

    # Totaux
    doc.add_heading("Recapitulatif", level=2)
    total_p = participants.get("total_presents", "")
    total_r = participants.get("total_representes", "")
    total_v = participants.get("total_votants", "")
    total_voix = participants.get("total_voix", "")
    quorum = participants.get("quorum_atteint")
    quorum_requis = participants.get("quorum_requis", "")

    table_tot = doc.add_table(rows=0, cols=2)
    table_tot.style = "Table Grid"
    for label, val in [
        ("Total presents", str(total_p) if total_p else ""),
        ("Total representes", str(total_r) if total_r else ""),
        ("Total votants", str(total_v) if total_v else ""),
        (f"Total {col_unite.lower()}", str(total_voix) if total_voix else ""),
        ("Quorum requis", quorum_requis or "Selon statuts"),
        ("Quorum atteint", "OUI" if quorum is True else "NON" if quorum is False else "A verifier"),
    ]:
        row = table_tot.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val

    doc.add_paragraph()

    # Certification
    p_cert = doc.add_paragraph()
    p_cert.add_run("Certifie exact par le president de seance :").bold = True
    doc.add_paragraph()
    doc.add_paragraph("Nom : ___________________________    Signature : ___________________________")
    doc.add_paragraph()

    # Disclaimer
    p_dis = doc.add_paragraph()
    p_dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_dis.add_run(
        f"Feuille de presence generee par AG Assistant — {datetime.now().strftime('%d/%m/%Y')} — Document provisoire"
    )
    run_d.italic = True
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(chemin_sortie)
    return chemin_sortie


def generer_feuille_presence_pdf(analyse: dict, chemin_sortie: str) -> str:
    """Genere une feuille de presence PDF."""
    infos = analyse.get("informations_generales", {})
    entite = _sanitiser(infos.get("entite", "Entite inconnue"))
    date_ag = _sanitiser(infos.get("date", "________________"))
    lieu = _sanitiser(infos.get("lieu", "________________"))
    type_asm = _sanitiser(infos.get("type_assemblee", "AG"))
    participants = analyse.get("participants", {})
    type_ag = analyse.get("type_ag", "autre")

    if type_ag == "copropriete":
        col_unite = "Milliemes"
    elif type_ag in ("pme_sas", "pme_sarl", "pme_autre"):
        col_unite = "Parts/Actions"
    else:
        col_unite = "Voix"

    pdf = FPDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    largeur = pdf.w - pdf.l_margin - pdf.r_margin

    # Titre
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(largeur, 8, "FEUILLE DE PRESENCE", align="C")
    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(largeur, 6, f"{type_asm.upper()} - {entite}", align="C")
    pdf.ln(4)

    # Infos
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(40, 6, "Entite :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 40, 6, entite, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(40, 6, "Date :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 40, 6, date_ag, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(40, 6, "Lieu :")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur - 40, 6, lieu[:80], new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Tableau presents
    pdf.set_font("Helvetica", "B", 10)
    pdf.multi_cell(largeur, 6, "MEMBRES PRESENTS")
    w_nom = largeur * 0.45
    w_voix = largeur * 0.2
    w_sig = largeur * 0.2
    w_obs = largeur - w_nom - w_voix - w_sig
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(w_nom, 6, "Nom / Prenom", border=1)
    pdf.cell(w_voix, 6, col_unite, border=1)
    pdf.cell(w_sig, 6, "Signature", border=1)
    pdf.cell(w_obs, 6, "Observations", border=1, new_x="LMARGIN", new_y="NEXT")

    presents = participants.get("presents", [])
    pdf.set_font("Helvetica", "", 9)
    if isinstance(presents, list):
        for p in presents:
            nom = _sanitiser(p.get("nom", "") if isinstance(p, dict) else str(p))[:50]
            voix = str(p.get("voix_ou_parts", "") if isinstance(p, dict) else "")
            pdf.cell(w_nom, 6, nom, border=1)
            pdf.cell(w_voix, 6, voix, border=1)
            pdf.cell(w_sig, 6, "", border=1)
            pdf.cell(w_obs, 6, "", border=1, new_x="LMARGIN", new_y="NEXT")
    for _ in range(max(0, 8 - (len(presents) if isinstance(presents, list) else 0))):
        pdf.cell(w_nom, 6, "", border=1)
        pdf.cell(w_voix, 6, "", border=1)
        pdf.cell(w_sig, 6, "", border=1)
        pdf.cell(w_obs, 6, "", border=1, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(3)

    # Recapitulatif quorum
    total_v = participants.get("total_votants", "")
    quorum = participants.get("quorum_atteint")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(largeur, 6, f"Total votants : {total_v}    Quorum atteint : {'OUI' if quorum is True else 'NON' if quorum is False else 'A verifier'}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(largeur, 6, "President de seance : _______________________________    Signature : ___________________", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(largeur, 5, f"Genere par AG Assistant le {datetime.now().strftime('%d/%m/%Y')} - Document provisoire", align="C")

    pdf.output(chemin_sortie)
    return chemin_sortie
